import io
import mimetypes
from typing import Optional, List
from pathlib import Path
from loguru import logger

try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False
    logger.warning("python-magic未安装或libmagic库缺失，无法根据内容检测文件类型")

# 文件内容提取相关导入
try:
    import pypdf
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logger.warning("PDF处理库未安装，无法提取PDF内容")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx未安装，无法提取Word文档内容")

try:
    import openpyxl
    import pandas as pd
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False
    logger.warning("Excel处理库未安装，无法提取Excel内容")

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    logger.warning("OCR库未安装，无法进行图片文字识别")

# 支持的文件类型
SUPPORTED_FILE_TYPES = {
    # 文档类型
    'application/pdf': ['.pdf'],
    'application/msword': ['.doc'],
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
    'application/vnd.ms-excel': ['.xls'],
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': ['.xlsx'],
    'text/plain': ['.txt'],
    'text/markdown': ['.md'],
    'text/csv': ['.csv'],
    'application/json': ['.json'],
    'application/xml': ['.xml'],
    'text/html': ['.html', '.htm'],
    
    # 图片类型
    'image/jpeg': ['.jpg', '.jpeg'],
    'image/png': ['.png'],
    'image/gif': ['.gif'],
    'image/bmp': ['.bmp'],
    'image/tiff': ['.tiff', '.tif'],
    'image/webp': ['.webp'],
    
    # 音频类型
    'audio/mpeg': ['.mp3'],
    'audio/wav': ['.wav'],
    'audio/ogg': ['.ogg'],
    'audio/flac': ['.flac'],
    
    # 视频类型
    'video/mp4': ['.mp4'],
    'video/avi': ['.avi'],
    'video/mov': ['.mov'],
    'video/wmv': ['.wmv'],
    'video/flv': ['.flv'],
    
    # 压缩文件
    'application/zip': ['.zip'],
    'application/x-rar-compressed': ['.rar'],
    'application/x-7z-compressed': ['.7z'],
    'application/x-tar': ['.tar'],
    'application/gzip': ['.gz'],
    
    # 代码文件
    'text/x-python': ['.py'],
    'text/javascript': ['.js'],
    'text/x-java-source': ['.java'],
    'text/x-c': ['.c'],
    'text/x-c++': ['.cpp', '.cxx'],
    'text/x-csharp': ['.cs'],
    'text/x-php': ['.php'],
    'text/x-ruby': ['.rb'],
    'text/x-go': ['.go'],
    'text/x-rust': ['.rs'],
    'text/x-sql': ['.sql'],
    'text/css': ['.css'],
    'text/x-scss': ['.scss'],
    'text/x-less': ['.less'],
}

# 文件大小限制（字节）
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
MAX_IMAGE_SIZE = 10 * 1024 * 1024   # 10MB
MAX_VIDEO_SIZE = 500 * 1024 * 1024  # 500MB

def get_file_type(filename: str) -> Optional[str]:
    """
    根据文件名获取MIME类型
    
    Args:
        filename: 文件名
        
    Returns:
        Optional[str]: MIME类型
    """
    try:
        mime_type, _ = mimetypes.guess_type(filename)
        return mime_type
    except Exception as e:
        logger.error(f"获取文件类型失败: {e}")
        return None

def get_file_type_by_content(file_content: bytes) -> Optional[str]:
    """
    根据文件内容获取MIME类型
    
    Args:
        file_content: 文件内容
        
    Returns:
        Optional[str]: MIME类型
    """
    if not HAS_MAGIC:
        return None
        
    try:
        mime_type = magic.from_buffer(file_content, mime=True)
        return mime_type
    except Exception as e:
        logger.error(f"根据内容获取文件类型失败: {e}")
        return None

def validate_file_type(filename: str) -> bool:
    """
    验证文件类型是否支持
    
    Args:
        filename: 文件名
        
    Returns:
        bool: 是否支持
    """
    try:
        file_extension = Path(filename).suffix.lower()
        
        # 检查扩展名是否在支持列表中
        for mime_type, extensions in SUPPORTED_FILE_TYPES.items():
            if file_extension in extensions:
                return True
        
        return False
        
    except Exception as e:
        logger.error(f"验证文件类型失败: {e}")
        return False

def validate_file_size(file_size: int, file_type: Optional[str] = None) -> bool:
    """
    验证文件大小是否符合限制
    
    Args:
        file_size: 文件大小（字节）
        file_type: 文件类型
        
    Returns:
        bool: 是否符合限制
    """
    try:
        if file_type:
            if file_type.startswith('image/'):
                return file_size <= MAX_IMAGE_SIZE
            elif file_type.startswith('video/'):
                return file_size <= MAX_VIDEO_SIZE
        
        return file_size <= MAX_FILE_SIZE
        
    except Exception as e:
        logger.error(f"验证文件大小失败: {e}")
        return False

def extract_text_from_pdf(file_stream: io.BytesIO) -> str:
    """
    从PDF文件提取文本
    
    Args:
        file_stream: PDF文件流
        
    Returns:
        str: 提取的文本内容
    """
    if not HAS_PDF:
        return ""
    
    try:
        text = ""
        
        # 首先尝试使用pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(file_stream) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber提取失败，尝试pypdf: {e}")
            
            # 回退到pypdf
            file_stream.seek(0)
            try:
                reader = pypdf.PdfReader(file_stream)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e2:
                logger.error(f"pypdf提取也失败: {e2}")
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"PDF文本提取失败: {e}")
        return ""

def extract_text_from_docx(file_stream: io.BytesIO) -> str:
    """
    从Word文档提取文本
    
    Args:
        file_stream: Word文档流
        
    Returns:
        str: 提取的文本内容
    """
    if not HAS_DOCX:
        return ""
    
    try:
        doc = Document(file_stream)
        text = ""
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Word文档文本提取失败: {e}")
        return ""

def extract_text_from_xlsx(file_stream: io.BytesIO) -> str:
    """
    从Excel文件提取文本
    
    Args:
        file_stream: Excel文件流
        
    Returns:
        str: 提取的文本内容
    """
    if not HAS_EXCEL:
        return ""
    
    try:
        # 使用openpyxl读取Excel
        workbook = openpyxl.load_workbook(file_stream)
        text = ""
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"工作表: {sheet_name}\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_text = []
                for cell in row:
                    if cell is not None:
                        row_text.append(str(cell))
                    else:
                        row_text.append("")
                text += "\t".join(row_text) + "\n"
            
            text += "\n"
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Excel文本提取失败: {e}")
        return ""

def extract_text_from_image(file_stream: io.BytesIO) -> str:
    """
    从图片提取文本（OCR）
    
    Args:
        file_stream: 图片文件流
        
    Returns:
        str: 提取的文本内容
    """
    if not HAS_OCR:
        return ""
    
    try:
        image = Image.open(file_stream)
        text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        return text.strip()
        
    except Exception as e:
        logger.error(f"图片OCR提取失败: {e}")
        return ""

def generate_file_hash(file_content: bytes) -> str:
    """
    生成文件哈希值
    
    Args:
        file_content: 文件内容
        
    Returns:
        str: 文件哈希值
    """
    try:
        import hashlib
        return hashlib.md5(file_content).hexdigest()
    except Exception as e:
        logger.error(f"生成文件哈希失败: {e}")
        return ""

def get_file_category(file_type: str) -> str:
    """
    根据文件类型获取文件分类
    
    Args:
        file_type: MIME类型
        
    Returns:
        str: 文件分类
    """
    try:
        if file_type.startswith('image/'):
            return 'image'
        elif file_type.startswith('video/'):
            return 'video'
        elif file_type.startswith('audio/'):
            return 'audio'
        elif file_type in ['application/pdf', 'application/msword', 
                          'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            return 'document'
        elif file_type in ['application/vnd.ms-excel', 
                          'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            return 'spreadsheet'
        elif file_type.startswith('text/'):
            return 'text'
        elif 'zip' in file_type or 'rar' in file_type or 'tar' in file_type:
            return 'archive'
        else:
            return 'other'
            
    except Exception as e:
        logger.error(f"获取文件分类失败: {e}")
        return 'other'

def format_file_size(size_bytes: int) -> str:
    """
    格式化文件大小
    
    Args:
        size_bytes: 文件大小（字节）
        
    Returns:
        str: 格式化后的文件大小
    """
    try:
        if size_bytes == 0:
            return "0 B"
        
        size_names = ["B", "KB", "MB", "GB", "TB"]
        i = 0
        size = float(size_bytes)
        
        while size >= 1024.0 and i < len(size_names) - 1:
            size /= 1024.0
            i += 1
        
        return f"{size:.1f} {size_names[i]}"
        
    except Exception as e:
        logger.error(f"格式化文件大小失败: {e}")
        return f"{size_bytes} B"

def is_safe_filename(filename: str) -> bool:
    """
    检查文件名是否安全
    
    Args:
        filename: 文件名
        
    Returns:
        bool: 是否安全
    """
    try:
        # 检查危险字符
        dangerous_chars = ['..', '/', '\\', ':', '*', '?', '"', '<', '>', '|']
        for char in dangerous_chars:
            if char in filename:
                return False
        
        # 检查文件名长度
        if len(filename) > 255:
            return False
        
        # 检查是否为空
        if not filename.strip():
            return False
        
        return True
        
    except Exception as e:
        logger.error(f"检查文件名安全性失败: {e}")
        return False

def sanitize_filename(filename: str) -> str:
    """
    清理文件名，使其安全
    
    Args:
        filename: 原始文件名
        
    Returns:
        str: 清理后的文件名
    """
    try:
        import re
        
        # 移除危险字符
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # 移除连续的点
        filename = re.sub(r'\.{2,}', '.', filename)
        
        # 限制长度
        if len(filename) > 255:
            name, ext = Path(filename).stem, Path(filename).suffix
            max_name_len = 255 - len(ext)
            filename = name[:max_name_len] + ext
        
        # 确保不为空
        if not filename.strip():
            filename = "unnamed_file"
        
        return filename
        
    except Exception as e:
        logger.error(f"清理文件名失败: {e}")
        return "unnamed_file" 